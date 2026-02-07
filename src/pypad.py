import tkinter as tk
from tkinter import *
from tkinter import filedialog, messagebox, font, colorchooser
from tkinter.ttk import Separator
import os
import platform

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

class EnhancedWordPad:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("PyPad - Enhanced Text Editor")
        self.root.geometry("1100x700")     
        self.current_file = None
        self.undo_stack = []
        self.redo_stack = []
        self.autosave_enabled = False
        self.dark_mode = False
        self.available_fonts = font.families()
        self.default_font = "Times New Roman" if "Times New Roman" in self.available_fonts else "Arial"
        self.current_font_family = self.default_font
        self.current_font_size = 12
        self.current_font_weight = "normal"
        self.current_font_slant = "roman"
        self.current_font_underline = False   
        if not HAVE_DOCX:
            print("Warning: python-docx library not installed. Install with: pip install python-docx")
        
        self.create_menubar()
        self.create_toolbar()
        self.create_text_area()
        self.create_statusbar()
        self.bind_shortcuts()      
        if self.autosave_enabled:
            self.root.after(300000, self.autosave)
        
    def create_menubar(self):
        menubar = Menu(self.root)
        self.root.config(menu=menubar)      
        file_menu = Menu(menubar, tearoff=0)
        file_menu.add_command(label="New", command=self.new_file, accelerator="Ctrl+N")
        file_menu.add_command(label="Open...", command=self.open_file, accelerator="Ctrl+O")
        file_menu.add_command(label="Save", command=self.save_file, accelerator="Ctrl+S")
        file_menu.add_command(label="Save As...", command=self.save_as, accelerator="Ctrl+Shift+S")
        file_menu.add_separator()
        file_menu.add_command(label="Preview Document...", command=self.preview_document)
        file_menu.add_command(label="Print...", command=self.print_file, accelerator="Ctrl+P")
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.exit_app, accelerator="Alt+F4")
        menubar.add_cascade(label="File", menu=file_menu)      
        edit_menu = Menu(menubar, tearoff=0)
        edit_menu.add_command(label="Undo", command=self.undo, accelerator="Ctrl+Z")
        edit_menu.add_command(label="Redo", command=self.redo, accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="Cut", command=self.cut, accelerator="Ctrl+X")
        edit_menu.add_command(label="Copy", command=self.copy, accelerator="Ctrl+C")
        edit_menu.add_command(label="Paste", command=self.paste, accelerator="Ctrl+V")
        edit_menu.add_command(label="Select All", command=self.select_all, accelerator="Ctrl+A")
        edit_menu.add_separator()
        edit_menu.add_command(label="Find...", command=self.find_text, accelerator="Ctrl+F")
        edit_menu.add_command(label="Replace...", command=self.replace_text, accelerator="Ctrl+H")
        menubar.add_cascade(label="Edit", menu=edit_menu)   
        format_menu = Menu(menubar, tearoff=0)
        format_menu.add_command(label="Font...", command=self.choose_font)
        format_menu.add_command(label="Text Color...", command=self.choose_text_color)
        format_menu.add_command(label="Background Color...", command=self.choose_bg_color)
        format_menu.add_separator()
        format_menu.add_checkbutton(label="Word Wrap", command=self.toggle_wrap)
        format_menu.add_checkbutton(label="Dark Mode", command=self.toggle_dark_mode)
        menubar.add_cascade(label="Format", menu=format_menu)      
        view_menu = Menu(menubar, tearoff=0)
        view_menu.add_checkbutton(label="Toolbar", command=self.toggle_toolbar)
        view_menu.add_checkbutton(label="Status Bar", command=self.toggle_statusbar)
        view_menu.add_separator()
        view_menu.add_command(label="Zoom In", command=self.zoom_in, accelerator="Ctrl++")
        view_menu.add_command(label="Zoom Out", command=self.zoom_out, accelerator="Ctrl+-")
        view_menu.add_command(label="Reset Zoom", command=self.reset_zoom, accelerator="Ctrl+0")
        menubar.add_cascade(label="View", menu=view_menu)     
        help_menu = Menu(menubar, tearoff=0)
        help_menu.add_command(label="About", command=self.show_about)
        help_menu.add_command(label="Shortcuts", command=self.show_shortcuts)
        menubar.add_cascade(label="Help", menu=help_menu)
    
    def create_toolbar(self):
        self.toolbar = Frame(self.root, bd=1, relief=RAISED)
        self.toolbar.pack(side=TOP, fill=X)     
        Button(self.toolbar, text="New", command=self.new_file).pack(side=LEFT, padx=2, pady=2)
        Button(self.toolbar, text="Open", command=self.open_file).pack(side=LEFT, padx=2, pady=2)
        Button(self.toolbar, text="Save", command=self.save_file).pack(side=LEFT, padx=2, pady=2)      
        Separator(self.toolbar, orient=VERTICAL).pack(side=LEFT, padx=5, pady=2)      
        self.bold_btn = Button(self.toolbar, text="B", font=("Arial", 10, "bold"), command=self.toggle_bold)
        self.bold_btn.pack(side=LEFT, padx=2, pady=2)      
        self.italic_btn = Button(self.toolbar, text="I", font=("Arial", 10, "italic"), command=self.toggle_italic)
        self.italic_btn.pack(side=LEFT, padx=2, pady=2)      
        self.underline_btn = Button(self.toolbar, text="U", font=("Arial", 10, "underline"), command=self.toggle_underline)
        self.underline_btn.pack(side=LEFT, padx=2, pady=2)
        self.font_family_var = StringVar(value=self.default_font)
        self.font_family_combo = tk.Listbox(self.toolbar, height=1, width=20, 
                                           exportselection=0, selectmode=SINGLE)
        self.font_family_combo.insert(0, self.default_font)
        self.font_family_combo.bind('<<ListboxSelect>>', self.on_font_family_select)
        self.font_family_combo.pack(side=LEFT, padx=2, pady=2)
        self.update_font_list()    
        self.font_size = Spinbox(self.toolbar, from_=8, to=72, width=5, command=self.change_font_size)
        self.font_size.delete(0, "end")
        self.font_size.insert(0, str(self.current_font_size))
        self.font_size.pack(side=LEFT, padx=2, pady=2)    
        Button(self.toolbar, text="◀", command=lambda: self.align_text("left")).pack(side=LEFT, padx=2, pady=2)
        Button(self.toolbar, text="►", command=lambda: self.align_text("center")).pack(side=LEFT, padx=2, pady=2)
        Button(self.toolbar, text="▶", command=lambda: self.align_text("right")).pack(side=LEFT, padx=2, pady=2)  
        self.word_count_label = Label(self.toolbar, text="Words: 0")
        self.word_count_label.pack(side=RIGHT, padx=10)
    
    def update_font_list(self):
        """Update the font family listbox with available fonts"""
        self.font_family_combo.delete(0, END)
        
        # Filter and sort fonts
        filtered_fonts = []
        for f in self.available_fonts:
            if f and len(f) > 0 and not f.startswith('@'):
                filtered_fonts.append(f)
        filtered_fonts.sort()
        
        # Add Times New Roman first if available
        if "Times New Roman" in filtered_fonts:
            filtered_fonts.remove("Times New Roman")
            filtered_fonts.insert(0, "Times New Roman")
        
        # Add other common fonts next
        preferred_order = ["Arial", "Courier New", "Georgia", "Verdana", 
                          "Tahoma", "Trebuchet MS", "Comic Sans MS"]
        for font_name in preferred_order:
            if font_name in filtered_fonts and font_name != "Times New Roman":
                filtered_fonts.remove(font_name)
                filtered_fonts.insert(1, font_name)
        
        # Add remaining fonts
        for f in filtered_fonts:
            self.font_family_combo.insert(END, f)
        
        # Select the current font
        try:
            index = filtered_fonts.index(self.current_font_family)
            self.font_family_combo.select_set(index)
            self.font_family_combo.see(index)
        except ValueError:
            if filtered_fonts:
                self.font_family_combo.select_set(0)
    
    def on_font_family_select(self, event=None):
        """Handle font family selection from combobox"""
        selection = self.font_family_combo.curselection()
        if selection:
            selected_font = self.font_family_combo.get(selection[0])
            self.current_font_family = selected_font
            self.apply_current_font()
    
    def create_text_area(self):
        main_frame = Frame(self.root)
        main_frame.pack(fill=BOTH, expand=True)      
        self.line_numbers = Text(main_frame, width=4, padx=3, takefocus=0, border=0, background='lightgrey', state='disabled')
        self.line_numbers.pack(side=LEFT, fill=Y)      
        # Set default font to Times New Roman if available, otherwise Arial
        initial_font = (self.default_font, self.current_font_size)
        self.text_area = Text(main_frame, wrap="word", undo=True, font=initial_font, selectbackground="lightblue")
        self.text_area.pack(side=LEFT, fill=BOTH, expand=True)      
        y_scrollbar = Scrollbar(self.text_area)
        y_scrollbar.pack(side=RIGHT, fill=Y)
        self.text_area.config(yscrollcommand=y_scrollbar.set)
        y_scrollbar.config(command=self.text_area.yview)      
        x_scrollbar = Scrollbar(main_frame, orient=HORIZONTAL)
        x_scrollbar.pack(side=BOTTOM, fill=X)
        self.text_area.config(xscrollcommand=x_scrollbar.set)
        x_scrollbar.config(command=self.text_area.xview)    
        self.text_area.bind('<KeyRelease>', self.update_word_count)
        self.text_area.bind('<Configure>', self.update_line_numbers)
        self.text_area.bind('<MouseWheel>', self.update_line_numbers)
    
    def create_statusbar(self):
        self.status_bar = Label(self.root, text=f"Ready | Font: {self.current_font_family}, {self.current_font_size}pt | Line: 1, Column: 1", bd=1, relief=SUNKEN, anchor=W)
        self.status_bar.pack(side=BOTTOM, fill=X)    
        self.text_area.bind('<KeyRelease>', self.update_cursor_position)
        self.text_area.bind('<ButtonRelease>', self.update_cursor_position)
    
    def bind_shortcuts(self):
        self.root.bind('<Control-n>', lambda e: self.new_file())
        self.root.bind('<Control-o>', lambda e: self.open_file())
        self.root.bind('<Control-s>', lambda e: self.save_file())
        self.root.bind('<Control-Shift-S>', lambda e: self.save_as())
        self.root.bind('<Control-f>', lambda e: self.find_text())
        self.root.bind('<Control-h>', lambda e: self.replace_text())
        self.root.bind('<Control-a>', lambda e: self.select_all())
        self.root.bind('<Control-b>', lambda e: self.toggle_bold())
        self.root.bind('<Control-i>', lambda e: self.toggle_italic())
        self.root.bind('<Control-u>', lambda e: self.toggle_underline())
        self.root.bind('<Control-plus>', lambda e: self.zoom_in())
        self.root.bind('<Control-minus>', lambda e: self.zoom_out())
        self.root.bind('<Control-0>', lambda e: self.reset_zoom())
    
    def apply_current_font(self):
        """Apply the current font settings to the text area"""
        font_elements = [self.current_font_family, self.current_font_size]
        if self.current_font_weight == "bold":
            font_elements.append("bold")
        if self.current_font_slant == "italic":
            font_elements.append("italic")
        if self.current_font_underline:
            font_elements.append("underline")
        
        font_tuple = tuple(font_elements)
        self.text_area.config(font=font_tuple)
        self.status_bar.config(text=f"Ready | Font: {self.current_font_family}, {self.current_font_size}pt | Line: 1, Column: 1")
        self.font_size.delete(0, END)
        self.font_size.insert(0, str(self.current_font_size))
    
    def new_file(self):
        if self.check_unsaved_changes():
            self.text_area.delete(1.0, END)
            self.current_file = None
            self.root.title("PyPad - New Document")
            self.update_word_count()
    
    def open_file(self):
        if self.check_unsaved_changes():
            file_path = filedialog.askopenfilename(
                filetypes=[
                    ("Text files", "*.txt"),
                    ("Python files", "*.py"),
                    ("HTML files", "*.html;*.htm"),
                    ("Word documents", "*.docx"),
                    ("All files", "*.*")
                ]
            )
            if file_path:
                self.load_file(file_path)
    
    def load_file(self, file_path):
        """Enhanced file loading with better .docx support"""
        try:
            if file_path.lower().endswith('.docx'):
                if not HAVE_DOCX:
                    messagebox.showerror("Error", 
                        "To open .docx files, please install python-docx:\n"
                        "pip install python-docx")
                    return
                
                # Enhanced .docx loading with formatting markers
                content = self.load_docx_with_formatting(file_path)
                self.text_area.delete(1.0, END)
                self.text_area.insert(1.0, content)
                
            else:
                # Try different encodings for text files
                encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
                content = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            content = file.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                if content is None:
                    # If all encodings fail, try binary mode as last resort
                    with open(file_path, 'rb') as file:
                        content = file.read().decode('utf-8', errors='replace')
                
                self.text_area.delete(1.0, END)
                self.text_area.insert(1.0, content)
            
            self.current_file = file_path
            self.root.title(f"PyPad - {os.path.basename(file_path)}")
            self.update_word_count()
            
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file: {str(e)}")
    
    def load_docx_with_formatting(self, file_path):
        """Enhanced .docx loader with formatting markers and content extraction"""
        try:
            doc = Document(file_path)
            content_lines = []
            
            # Document metadata
            content_lines.append("=" * 70)
            content_lines.append("DOCUMENT EXTRACTED FROM .DOCX FILE")
            content_lines.append("=" * 70)
            content_lines.append("")
            
            # Extract core properties if available
            core_props = doc.core_properties
            if core_props.author or core_props.created:
                content_lines.append("[Document Properties]")
                if core_props.author:
                    content_lines.append(f"Author: {core_props.author}")
                if core_props.created:
                    content_lines.append(f"Created: {core_props.created}")
                if core_props.title:
                    content_lines.append(f"Title: {core_props.title}")
                content_lines.append("")
            
            # Process paragraphs with formatting
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    if paragraph.runs:  # Empty paragraph with formatting
                        content_lines.append("[Empty paragraph]")
                    continue
                
                # Check paragraph style
                style_name = paragraph.style.name if paragraph.style else ""
                if 'Heading' in style_name:
                    heading_level = 1
                    if 'Heading 1' in style_name:
                        heading_level = 1
                    elif 'Heading 2' in style_name:
                        heading_level = 2
                    elif 'Heading 3' in style_name:
                        heading_level = 3
                    content_lines.append(f"\n{'#' * heading_level} {paragraph.text}")
                    continue
                
                # Process runs in paragraph
                line_text = ""
                for run in paragraph.runs:
                    run_text = run.text
                    if not run_text.strip():
                        line_text += run_text
                        continue
                    
                    # Apply formatting markers
                    formatted_text = run_text
                    
                    # Mark bold text
                    if run.bold:
                        formatted_text = f"**{formatted_text}**"
                    
                    # Mark italic text
                    if run.italic:
                        formatted_text = f"*{formatted_text}*"
                    
                    # Mark underline text
                    if run.underline:
                        formatted_text = f"_{formatted_text}_"
                    
                    # Mark highlighted text (if we can detect it)
                    if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                        formatted_text = f"[HIGHLIGHT]{formatted_text}[/HIGHLIGHT]"
                    
                    line_text += formatted_text
                
                if line_text.strip():
                    content_lines.append(line_text)
            
            # Process tables
            if doc.tables:
                content_lines.append("\n" + "=" * 70)
                content_lines.append("TABLES")
                content_lines.append("=" * 70)
                
                for table_idx, table in enumerate(doc.tables):
                    content_lines.append(f"\n[Table {table_idx + 1}]")
                    
                    # Get table dimensions
                    row_count = len(table.rows)
                    col_count = len(table.columns) if hasattr(table, 'columns') else 0
                    content_lines.append(f"Dimensions: {row_count} rows × {col_count} columns")
                    
                    # Extract table content
                    for row_idx, row in enumerate(table.rows):
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                # Limit cell text length for display
                                if len(cell_text) > 50:
                                    cell_text = cell_text[:47] + "..."
                                row_cells.append(cell_text)
                            else:
                                row_cells.append("[empty]")
                        
                        if row_cells:
                            content_lines.append(" | ".join(row_cells))
                    
                    content_lines.append("")
            
            # Footer with document info
            content_lines.append("\n" + "=" * 70)
            content_lines.append("DOCUMENT INFORMATION")
            content_lines.append("=" * 70)
            content_lines.append(f"Total paragraphs: {len(doc.paragraphs)}")
            content_lines.append(f"Total tables: {len(doc.tables)}")
            content_lines.append(f"Total sections: {len(doc.sections)}")
            
            # Check for images and other objects
            try:
                # Count inline shapes (images, charts, etc.)
                inline_shapes = []
                for section in doc.sections:
                    for paragraph in doc.paragraphs:
                        for inline in paragraph.runs:
                            if hasattr(inline, 'element') and inline.element.xpath('.//wp:inline'):
                                inline_shapes.append(inline)
                
                if inline_shapes:
                    content_lines.append(f"Images/Objects: {len(inline_shapes)} (not displayed in text view)")
            except:
                content_lines.append("Images/Objects: (unable to detect)")
            
            content_lines.append("\nNote: Formatting markers: **bold**, *italic*, _underline_")
            content_lines.append("=" * 70)
            
            return "\n".join(content_lines)
            
        except Exception as e:
            # Fallback to simple extraction
            try:
                doc = Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                content += "\n\n[Note: Using basic extraction due to error]"
                return content
            except:
                return f"Error loading .docx file: {str(e)}"
    
    def preview_document(self):
        """Show a preview window with document statistics"""
        if not self.current_file or not self.current_file.lower().endswith('.docx'):
            messagebox.showinfo("Preview", "Preview is only available for .docx files.")
            return
        
        if not HAVE_DOCX:
            messagebox.showerror("Error", "python-docx library not installed.")
            return
        
        try:
            doc = Document(self.current_file)
            
            # Create preview window
            preview_window = Toplevel(self.root)
            preview_window.title(f"Document Preview - {os.path.basename(self.current_file)}")
            preview_window.geometry("600x500")
            
            # Create text widget for preview
            preview_text = Text(preview_window, wrap="word", font=("Arial", 10))
            preview_text.pack(fill=BOTH, expand=True, padx=10, pady=10)
            
            # Add scrollbar
            scrollbar = Scrollbar(preview_text)
            scrollbar.pack(side=RIGHT, fill=Y)
            preview_text.config(yscrollcommand=scrollbar.set)
            scrollbar.config(command=preview_text.yview)
            
            # Add preview content
            preview_content = []
            preview_content.append("=" * 60)
            preview_content.append("DOCUMENT PREVIEW")
            preview_content.append("=" * 60)
            preview_content.append("")
            
            # Document statistics
            preview_content.append("STATISTICS:")
            preview_content.append("-" * 40)
            preview_content.append(f"Paragraphs: {len(doc.paragraphs)}")
            preview_content.append(f"Tables: {len(doc.tables)}")
            preview_content.append(f"Sections: {len(doc.sections)}")
            
            # Core properties
            core_props = doc.core_properties
            preview_content.append("\nPROPERTIES:")
            preview_content.append("-" * 40)
            if core_props.author:
                preview_content.append(f"Author: {core_props.author}")
            if core_props.title:
                preview_content.append(f"Title: {core_props.title}")
            if core_props.subject:
                preview_content.append(f"Subject: {core_props.subject}")
            if core_props.created:
                preview_content.append(f"Created: {core_props.created}")
            if core_props.modified:
                preview_content.append(f"Modified: {core_props.modified}")
            
            # Sample content
            preview_content.append("\nSAMPLE CONTENT (first 5 paragraphs):")
            preview_content.append("-" * 40)
            for i, para in enumerate(doc.paragraphs[:5]):
                if para.text.strip():
                    text = para.text
                    if len(text) > 100:
                        text = text[:97] + "..."
                    preview_content.append(f"{i+1}. {text}")
            
            # Tables info
            if doc.tables:
                preview_content.append("\nTABLE INFORMATION:")
                preview_content.append("-" * 40)
                for i, table in enumerate(doc.tables[:3]):
                    row_count = len(table.rows)
                    col_count = len(table.columns) if hasattr(table, 'columns') else 0
                    preview_content.append(f"Table {i+1}: {row_count} rows × {col_count} columns")
            
            preview_content.append("\n" + "=" * 60)
            preview_content.append("Note: This is a preview. Open the file to see full content.")
            preview_content.append("=" * 60)
            
            preview_text.insert(1.0, "\n".join(preview_content))
            preview_text.config(state='disabled')
            
        except Exception as e:
            messagebox.showerror("Preview Error", f"Could not generate preview: {str(e)}")
    
    def save_file(self):
        if self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_as()
    
    def save_as(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text files", "*.txt"),
                ("Python files", "*.py"),
                ("HTML files", "*.html"),
                ("Word documents", "*.docx"),
                ("All files", "*.*")
            ]
        )
        if file_path:
            self.current_file = file_path
            self.save_to_file(file_path)
            self.root.title(f"PyPad - {os.path.basename(file_path)}")
    
    def save_to_file(self, file_path):
        """Save content to a file based on its extension"""
        try:
            content = self.text_area.get(1.0, END)
            
            if file_path.lower().endswith('.docx'):
                if HAVE_DOCX:
                    doc = Document()
                    
                    # Try to preserve formatting markers
                    lines = content.split('\n')
                    for line in lines:
                        if line.strip():
                            para = doc.add_paragraph()
                            
                            # Simple formatting preservation
                            text = line
                            # Remove markers and apply formatting
                            if '**' in text:
                                # Handle bold markers
                                parts = text.split('**')
                                for i, part in enumerate(parts):
                                    if i % 2 == 0:  # Regular text
                                        run = para.add_run(part)
                                    else:  # Bold text
                                        run = para.add_run(part)
                                        run.bold = True
                            else:
                                para.add_run(text)
                    
                    doc.save(file_path)
                else:
                    messagebox.showerror("Error", 
                        "To save .docx files, please install python-docx:\n"
                        "pip install python-docx")
                    return
            else:
                # For text files, try UTF-8 first, fall back to system encoding
                try:
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(content)
                except:
                    with open(file_path, 'w') as file:
                        file.write(content)
            
            self.status_bar.config(text="File saved successfully")
            
        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {str(e)}")
    
    def print_file(self):
        if platform.system() == "Windows":
            content = self.text_area.get(1.0, END)
            temp_file = "temp_print.txt"
            try:
                with open(temp_file, 'w', encoding='utf-8') as f:
                    f.write(content)
                os.startfile(temp_file, "print")
            except Exception as e:
                messagebox.showerror("Error", f"Could not print: {str(e)}")
        else:
            messagebox.showinfo("Print", "Print functionality is limited on this platform")
    
    def check_unsaved_changes(self):
        content = self.text_area.get(1.0, END).strip()
        if content and self.current_file is None:
            response = messagebox.askyesnocancel("Save Changes", "Do you want to save changes?")
            if response is None:
                return False
            elif response:
                self.save_as()
        return True
    
    def undo(self):
        try:
            self.text_area.edit_undo()
        except:
            pass
    
    def redo(self):
        try:
            self.text_area.edit_redo()
        except:
            pass
    
    def cut(self):
        self.text_area.event_generate("<<Cut>>")
    
    def copy(self):
        self.text_area.event_generate("<<Copy>>")
    
    def paste(self):
        self.text_area.event_generate("<<Paste>>")
    
    def select_all(self):
        self.text_area.tag_add(SEL, "1.0", END)
        self.text_area.mark_set(INSERT, "1.0")
        self.text_area.see(INSERT)
    
    def find_text(self):
        find_window = Toplevel(self.root)
        find_window.title("Find")
        find_window.geometry("400x150")
        
        Label(find_window, text="Find:").pack(pady=5)
        find_entry = Entry(find_window, width=40)
        find_entry.pack(pady=5)
        
        def find_next():
            text_to_find = find_entry.get()
            if text_to_find:
                start_pos = self.text_area.search(text_to_find, INSERT, END)
                if start_pos:
                    end_pos = f"{start_pos}+{len(text_to_find)}c"
                    self.text_area.tag_remove(SEL, "1.0", END)
                    self.text_area.tag_add(SEL, start_pos, end_pos)
                    self.text_area.mark_set(INSERT, end_pos)
                    self.text_area.see(INSERT)
                else:
                    messagebox.showinfo("Find", "Text not found")
        
        Button(find_window, text="Find Next", command=find_next).pack(pady=10)
    
    def replace_text(self):
        replace_window = Toplevel(self.root)
        replace_window.title("Replace")
        replace_window.geometry("400x200")
        
        Label(replace_window, text="Find:").pack(pady=5)
        find_entry = Entry(replace_window, width=40)
        find_entry.pack(pady=5)
        
        Label(replace_window, text="Replace with:").pack(pady=5)
        replace_entry = Entry(replace_window, width=40)
        replace_entry.pack(pady=5)
        
        def replace_all():
            find_text = find_entry.get()
            replace_text = replace_entry.get()
            if find_text:
                content = self.text_area.get(1.0, END)
                new_content = content.replace(find_text, replace_text)
                self.text_area.delete(1.0, END)
                self.text_area.insert(1.0, new_content)
        
        Button(replace_window, text="Replace All", command=replace_all).pack(pady=10)
    
    def choose_font(self):
        font_window = Toplevel(self.root)
        font_window.title("Font Selection")
        font_window.geometry("500x500")
        font_family_frame = Frame(font_window)
        font_family_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)    
        Label(font_family_frame, text="Font Family:").pack(anchor=W)
        font_list_frame = Frame(font_family_frame)
        font_list_frame.pack(fill=BOTH, expand=True)
        scrollbar = Scrollbar(font_list_frame)
        scrollbar.pack(side=RIGHT, fill=Y)
        listbox = Listbox(font_list_frame, height=15, yscrollcommand=scrollbar.set, exportselection=0)
        listbox.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        filtered_fonts = []
        for f in self.available_fonts:
            if f and len(f) > 0 and not f.startswith('@'):
                filtered_fonts.append(f)
        filtered_fonts.sort()
        if "Times New Roman" in filtered_fonts:
            filtered_fonts.remove("Times New Roman")
            filtered_fonts.insert(0, "Times New Roman")
        
        for family in filtered_fonts:
            listbox.insert(END, family)

        try:
            index = filtered_fonts.index(self.current_font_family)
            listbox.select_set(index)
            listbox.see(index)
        except ValueError:
            if filtered_fonts:
                listbox.select_set(0)
        
        size_frame = Frame(font_window)
        size_frame.pack(fill=X, padx=10, pady=5)
        Label(size_frame, text="Size:").pack(side=LEFT, padx=5)
        font_size_spinbox = Spinbox(size_frame, from_=8, to=72, width=10)
        font_size_spinbox.delete(0, END)
        font_size_spinbox.insert(0, str(self.current_font_size))
        font_size_spinbox.pack(side=LEFT, padx=5)
        style_frame = Frame(font_window)
        style_frame.pack(fill=X, padx=10, pady=5)
        bold_var = IntVar(value=1 if self.current_font_weight == "bold" else 0)
        italic_var = IntVar(value=1 if self.current_font_slant == "italic" else 0)
        underline_var = IntVar(value=1 if self.current_font_underline else 0)
        Checkbutton(style_frame, text="Bold", variable=bold_var).pack(side=LEFT, padx=5)
        Checkbutton(style_frame, text="Italic", variable=italic_var).pack(side=LEFT, padx=5)
        Checkbutton(style_frame, text="Underline", variable=underline_var).pack(side=LEFT, padx=5)
        
        def apply_font():
            selection = listbox.curselection()
            if selection:
                selected_font = listbox.get(selection[0])
                self.current_font_family = selected_font
                
                try:
                    self.current_font_size = int(font_size_spinbox.get())
                except ValueError:
                    self.current_font_size = 12
                
                self.current_font_weight = "bold" if bold_var.get() else "normal"
                self.current_font_slant = "italic" if italic_var.get() else "roman"
                self.current_font_underline = bool(underline_var.get())
                self.apply_current_font()
                font_window.destroy()
        
        Button(font_window, text="Apply", command=apply_font).pack(pady=10)
        
        listbox.bind('<Double-Button-1>', lambda e: apply_font())
    
    def choose_text_color(self):
        color = colorchooser.askcolor(title="Choose text color")
        if color[1]:
            self.text_area.config(fg=color[1])
    
    def choose_bg_color(self):
        color = colorchooser.askcolor(title="Choose background color")
        if color[1]:
            self.text_area.config(bg=color[1])
            self.line_numbers.config(bg='lightgrey' if not self.dark_mode else 'gray30')
    
    def toggle_wrap(self):
        current_wrap = self.text_area.cget("wrap")
        new_wrap = "none" if current_wrap == "word" else "word"
        self.text_area.config(wrap=new_wrap)
    
    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.text_area.config(bg='gray20', fg='white', insertbackground='white')
            self.root.config(bg='gray15')
        else:
            self.text_area.config(bg='white', fg='black', insertbackground='black')
            self.root.config(bg='SystemButtonFace')
    
    def toggle_toolbar(self):
        if self.toolbar.winfo_ismapped():
            self.toolbar.pack_forget()
        else:
            self.toolbar.pack(side=TOP, fill=X)
    
    def toggle_statusbar(self):
        if self.status_bar.winfo_ismapped():
            self.status_bar.pack_forget()
        else:
            self.status_bar.pack(side=BOTTOM, fill=X)
    
    def zoom_in(self):
        self.current_font_size += 1
        self.apply_current_font()
    
    def zoom_out(self):
        if self.current_font_size > 6:
            self.current_font_size -= 1
            self.apply_current_font()
    
    def reset_zoom(self):
        self.current_font_family = self.default_font
        self.current_font_size = 12
        self.current_font_weight = "normal"
        self.current_font_slant = "roman"
        self.current_font_underline = False
        self.apply_current_font()
    
    def toggle_bold(self):
        self.current_font_weight = "bold" if self.current_font_weight != "bold" else "normal"
        self.apply_current_font()
    
    def toggle_italic(self):
        self.current_font_slant = "italic" if self.current_font_slant != "italic" else "roman"
        self.apply_current_font()
    
    def toggle_underline(self):
        self.current_font_underline = not self.current_font_underline
        self.apply_current_font()
    
    def change_font_size(self):
        try:
            new_size = int(self.font_size.get())
            self.current_font_size = new_size
            self.apply_current_font()
        except ValueError:
            pass
    
    def align_text(self, alignment):
        self.text_area.tag_configure(alignment, justify=alignment)
        try:
            self.text_area.tag_add(alignment, SEL_FIRST, SEL_LAST)
        except:
            pass
    
    def update_word_count(self, event=None):
        content = self.text_area.get(1.0, END).strip()
        words = len(content.split()) if content else 0
        self.word_count_label.config(text=f"Words: {words}")
    
    def update_cursor_position(self, event=None):
        cursor_pos = self.text_area.index(INSERT)
        line, col = cursor_pos.split('.')
        self.status_bar.config(text=f"Ready | Font: {self.current_font_family}, {self.current_font_size}pt | Line: {line}, Column: {int(col)+1}")
    
    def update_line_numbers(self, event=None):
        self.line_numbers.config(state=NORMAL)
        self.line_numbers.delete(1.0, END)
        
        line_count = self.text_area.index('end-1c').split('.')[0]
        for i in range(1, int(line_count)+1):
            self.line_numbers.insert(END, f"{i}\n")
        
        self.line_numbers.config(state=DISABLED)
    
    def autosave(self):
        if self.autosave_enabled and self.text_area.get(1.0, END).strip():
            temp_file = "autosave.txt"
            with open(temp_file, 'w') as f:
                f.write(self.text_area.get(1.0, END))
            self.root.after(300000, self.autosave)
    
    def show_about(self):
        about_text = """PyPad - Enhanced Text Editor with .docx Support
        
Version: 1.2
Created with Python 3.12 and Tkinter
        
Features:
• Rich text editing with Times New Roman as default font
• Multiple file formats (TXT, PY, HTML, DOCX)
• Enhanced .docx support with formatting markers
• Document preview and statistics
• Find and Replace
• Word count
• Dark mode
• Customizable interface
        
© 2026 PyPad - Your custom WordPad alternative, by Netcat"""
        messagebox.showinfo("About PyPad", about_text)
    
    def show_shortcuts(self):
        shortcuts = """Keyboard Shortcuts:
        
Ctrl+N: New file
Ctrl+O: Open file
Ctrl+S: Save file
Ctrl+Shift+S: Save As
Ctrl+P: Print
        
Ctrl+Z: Undo
Ctrl+Y: Redo
Ctrl+X: Cut
Ctrl+C: Copy
Ctrl+V: Paste
Ctrl+A: Select All
        
Ctrl+F: Find
Ctrl+H: Replace
        
Ctrl+B: Bold
Ctrl+I: Italic
Ctrl+U: Underline
        
Ctrl++: Zoom In
Ctrl+-: Zoom Out
Ctrl+0: Reset Zoom
        
Alt+F4: Exit"""
        messagebox.showinfo("Keyboard Shortcuts", shortcuts)
    
    def exit_app(self):
        if self.check_unsaved_changes():
            self.root.quit()
    
    def run(self):
        self.root.protocol("WM_DELETE_WINDOW", self.exit_app)
        self.root.mainloop()

if __name__ == "__main__":
    app = EnhancedWordPad()
    app.run()
